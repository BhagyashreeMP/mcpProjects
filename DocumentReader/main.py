from mcp.server.fastmcp import FastMCP
from docx import Document
import os

# Initialize MCP server
mcp = FastMCP("WordSearchSummarizer")

# --- Mock Word file database ---
mock_doc_db = {
    "DOC001": {
        "title": "Project Proposal",
        "file_path": r"C:\Users\User\OneDrive\Documents\JSW training report.docx"  # fixed path
    }
    # Add more documents as needed
}

# --- Tools ---

@mcp.tool()
def read_document(document_id: str) -> str:
    """Read full text of a Word document by ID"""
    data = mock_doc_db.get(document_id)
    if not data:
        return f"Document ID '{document_id}' not found."
    if not os.path.exists(data["file_path"]):
        return f"File '{data['file_path']}' not found."

    doc = Document(data["file_path"])
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return text or "Document is empty."

@mcp.tool()
def search_keyword(keyword: str) -> str:
    """Search keyword across all documents and return snippets"""
    results = []
    for doc_id, info in mock_doc_db.items():
        if not os.path.exists(info["file_path"]):
            continue
        doc = Document(info["file_path"])
        for para in doc.paragraphs:
            if keyword.lower() in para.text.lower():
                snippet = para.text.strip()
                results.append(f"[{doc_id} - {info['title']}] {snippet}")
    return "\n\n".join(results) if results else f"No results found for '{keyword}'."

@mcp.tool()
def summarize_text(text: str, style: str = "concise") -> str:
    """Generate AI summary of given text"""
    summary_length = 200  # characters
    if style == "concise":
        return text[:summary_length] + ("..." if len(text) > summary_length else "")
    elif style == "detailed":
        return text[:summary_length*2] + ("..." if len(text) > summary_length*2 else "")
    elif style == "bullet":
        lines = text.split(". ")
        bullets = "\n".join([f"- {line.strip()}" for line in lines[:5]])
        return bullets
    return text[:summary_length]

# --- Resources ---

@mcp.resource("doc://{document_id}")
def get_document(document_id: str) -> str:
    return read_document(document_id)

@mcp.resource("search://{keyword}")
def search_resource(keyword: str) -> str:
    return search_keyword(keyword)

# Make 'style' optional so it doesn't break FastMCP URI matching
@mcp.resource("summarize://{document_id}")
def summarize_document(document_id):
    """Summarize a document with optional style passed as query param"""
    from mcp.server.fastmcp import request  # FastMCP exposes request context

    # Get optional style parameter from query, default to 'concise'
    style = request.query.get("style", "concise")

    text = read_document(document_id)
    if text.startswith("Document ID") or text.startswith("File"):
        return text
    return summarize_text(text, style)


@mcp.resource("summarize_keyword://{keyword}")
def summarize_keyword_results(keyword):
    """Summarize keyword search results with optional style"""
    from mcp.server.fastmcp import request

    # Get optional style parameter from query, default to 'concise'
    style = request.query.get("style", "concise")

    results = search_keyword(keyword)
    if results.startswith("No results"):
        return results
    return summarize_text(results, style)


# --- Prompts ---

@mcp.prompt()
def greet_user(name: str, style: str = "friendly") -> str:
    greetings = {
        "friendly": f"Hi {name}! Welcome to the Word Search & Summarizer.",
        "formal": f"Hello {name}, welcome to the Word Search & Summarizer system.",
        "casual": f"Hey {name}! Ready to search some docs?",
    }
    return greetings.get(style, greetings["friendly"])

@mcp.prompt()
def summarize_prompt(text: str, style: str = "concise") -> str:
    return summarize_text(text, style)

# --- Run MCP Server ---
if __name__ == "__main__":
    print("Starting Word Search & Summarizer MCP server...")
    mcp.run()
